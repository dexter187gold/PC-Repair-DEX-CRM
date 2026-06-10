import sys
import json
import os
from datetime import datetime
from kivy.app import App
from kivy.uix.screenmanager import ScreenManager, Screen
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.label import Label
from kivy.uix.textinput import TextInput
from kivy.uix.button import Button
from kivy.uix.spinner import Spinner
from kivy.uix.checkbox import CheckBox
from kivy.uix.scrollview import ScrollView
from kivy.core.window import Window
from kivy.uix.floatlayout import FloatLayout
from kivymd.app import MDApp
from kivymd.uix.button import MDRaisedButton
from kivymd.uix.textfield import MDTextField
from kivymd.uix.label import MDLabel
from kivymd.uix.toolbar import MDTopAppBar
from kivymd.uix.menu import MDDropdownMenu
import docx
from docx.shared import Inches
from PIL import Image
import io

class Client:
    def __init__(self, name, phone, email, address):
        self.name = name
        self.phone = phone
        self.email = email
        self.address = address

class MainApp(MDApp):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.theme_cls.theme_style = "Light"
        self.theme_cls.primary_palette = "Blue"
        self.clients = self.load_clients()
        self.current_job_number = self.load_job_number()
        self.before_photo = None
        self.after_photo = None

    def build(self):
        self.sm = ScreenManager()
        self.main_screen = MainScreen(name='main', app=self)
        self.sm.add_widget(self.main_screen)
        return self.sm

    def load_clients(self):
        if os.path.exists('clients.json'):
            with open('clients.json', 'r') as f:
                return json.load(f)
        return []

    def save_clients(self):
        with open('clients.json', 'w') as f:
            json.dump(self.clients, f)

    def load_job_number(self):
        if os.path.exists('job_number.json'):
            with open('job_number.json', 'r') as f:
                data = json.load(f)
                return data.get('number', 1)
        return 1

    def save_job_number(self):
        with open('job_number.json', 'w') as f:
            json.dump({'number': self.current_job_number}, f)

    def toggle_theme(self):
        self.theme_cls.theme_style = "Dark" if self.theme_cls.theme_style == "Light" else "Light"

class MainScreen(Screen):
    def __init__(self, app, **kwargs):
        super().__init__(**kwargs)
        self.app = app
        layout = BoxLayout(orientation='vertical', padding=10, spacing=10)

        # Top Bar
        toolbar = MDTopAppBar(title="PC Repair DEX CRM V1", md_bg_color=[0, 0.5, 1, 1])
        layout.add_widget(toolbar)

        # Theme Toggle
        theme_btn = MDRaisedButton(text="Toggle Dark/Light", on_release=lambda x: self.app.toggle_theme())
        layout.add_widget(theme_btn)

        # Client Section
        client_layout = BoxLayout(orientation='vertical', size_hint_y=None, height=200)
        client_layout.add_widget(MDLabel(text="Client Info", font_style="H5"))
        
        self.client_name = MDTextField(hint_text="Client / Business Name")
        self.client_phone = MDTextField(hint_text="Phone")
        self.client_email = MDTextField(hint_text="Email")
        self.client_address = MDTextField(hint_text="Address")
        
        client_layout.add_widget(self.client_name)
        client_layout.add_widget(self.client_phone)
        client_layout.add_widget(self.client_email)
        client_layout.add_widget(self.client_address)
        layout.add_widget(client_layout)

        # Tier Selection
        self.tier_spinner = Spinner(text='Enterprise', values=('Small', 'Medium', 'Enterprise'))
        layout.add_widget(self.tier_spinner)

        # Job Details
        job_layout = BoxLayout(orientation='vertical', size_hint_y=None, height=150)
        self.technician = MDTextField(hint_text="Technician Name")
        self.issue = MDTextField(hint_text="Issue Reported")
        job_layout.add_widget(self.technician)
        job_layout.add_widget(self.issue)
        layout.add_widget(job_layout)

        # Photo Buttons
        photo_layout = BoxLayout(orientation='horizontal', size_hint_y=None, height=50)
        before_btn = MDRaisedButton(text="Before Photo (Optional)", on_release=self.select_before_photo)
        after_btn = MDRaisedButton(text="After Photo (Optional)", on_release=self.select_after_photo)
        photo_layout.add_widget(before_btn)
        photo_layout.add_widget(after_btn)
        layout.add_widget(photo_layout)

        # Generate Button
        generate_btn = MDRaisedButton(text="Generate Report", on_release=self.generate_report)
        layout.add_widget(generate_btn)

        scroll = ScrollView()
        scroll.add_widget(layout)
        self.add_widget(scroll)

    def select_before_photo(self, instance):
        # Simplified - in real app use file chooser
        self.app.before_photo = "before_photo.jpg"  # Placeholder
        print("Before photo selected (placeholder)")

    def select_after_photo(self, instance):
        self.app.after_photo = "after_photo.jpg"  # Placeholder
        print("After photo selected (placeholder)")

    def generate_report(self, instance):
        tier = self.tier_spinner.text[0]
        year = datetime.now().year
        month = f"{datetime.now().month:02d}"
        job_id = f"DEX-{tier}-{year}-{month}-{self.app.current_job_number:04d}"
        
        # Create docx
        doc = docx.Document()
        doc.add_heading(f'PC Repair DEX - {tier} Business Service Report', 0)
        doc.add_paragraph(f'Report ID: {job_id}')
        doc.add_paragraph(f'Date: {datetime.now().strftime("%d/%m/%Y")}')
        
        # Add client info
        doc.add_heading('Client Information', level=1)
        doc.add_paragraph(f"Name: {self.client_name.text}")
        doc.add_paragraph(f"Phone: {self.client_phone.text}")
        # ... add more fields

        # Save
        filename = f"Report_{job_id}.docx"
        os.makedirs("Generated_Reports", exist_ok=True)
        doc.save(os.path.join("Generated_Reports", filename))
        
        self.app.current_job_number += 1
        self.app.save_job_number()
        
        print(f"Report generated: {filename}")

if __name__ == '__main__':
    MainApp().run()
